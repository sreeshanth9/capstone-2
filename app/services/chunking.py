# import re
# import nltk
# from nltk.tokenize import sent_tokenize
# from PyPDF2 import PdfReader
# from docx import Document

# nltk.download('punkt')

# class ChunkingService:
#     def __init__(self, chunk_size=200, overlap=50):
#         self.chunk_size = chunk_size
#         self.overlap = overlap

#     def extract_text(self, file_path):
#         text = ""
#         try:
#             if file_path.lower().endswith(".pdf"):
#                 with open(file_path, "rb") as file:
#                     reader = PdfReader(file)
#                     for page in reader.pages:
#                         text += page.extract_text() + "\n"
            
#             elif file_path.lower().endswith(".txt"):
#                 with open(file_path, "r", encoding="utf-8") as file:
#                     text = file.read().strip()

#             elif file_path.lower().endswith(".docx"):
#                 doc = Document(file_path)
#                 text = "\n".join([para.text for para in doc.paragraphs]).strip()

#             else:
#                 raise ValueError("Unsupported file format. Only PDF, TXT, and DOCX are allowed.")

#         except Exception as e:
#             print(f"Error extracting text from {file_path}: {e}")
        
#         return text

#     def clean_text(self, text):
#         text = re.sub(r'\s+', ' ', text)
#         return text.strip()

#     def split_into_sentences(self, text):
#         return sent_tokenize(text)

#     def recursive_chunking(self, sentences):
#         chunks = []
#         current_chunk = ""
        
#         for sentence in sentences:
#             if len(current_chunk) + len(sentence) <= self.chunk_size:
#                 current_chunk += " " + sentence
#             else:
#                 chunks.append(current_chunk.strip())
#                 current_chunk = sentence

#         if current_chunk:
#             chunks.append(current_chunk.strip())
        
#         overlapped_chunks = []
#         for i in range(len(chunks)):
#             start = max(0, i - 1)
#             overlap_text = " ".join(chunks[start:i + 1])
#             overlapped_chunks.append(overlap_text[:self.chunk_size].strip())
        
#         return overlapped_chunks

#     def process(self, pdf_path):
#         raw_text = self.extract_text(pdf_path)
#         cleaned_text = self.clean_text(raw_text)
#         sentences = self.split_into_sentences(cleaned_text)
#         return self.recursive_chunking(sentences)

# # if __name__ == "__main__":
# #     chunker = ChunkingService(chunk_size=300, overlap=50)
# #     chunks = chunker.process("sample.pdf")
# #     for i, chunk in enumerate(chunks):
# #         print(f"Chunk {i+1}:\n{chunk}\n")

# import os, json

# if __name__ == "__main__":
#     chunker = ChunkingService(chunk_size=300, overlap=50)
#     file_path = r"C:\Users\srees\OneDrive\Desktop\Semester-07\Report-Jury (27-11-2024).pdf"

#     chunks = chunker.process(file_path)

#     # Define the directory and file path
#     output_dir = r"C:\Users\srees\OneDrive\Desktop\Semester-08\capstone2-impl\app"
#     output_file = os.path.join(output_dir, "chunks1.json")

#     # Ensure the directory exists
#     os.makedirs(output_dir, exist_ok=True)

#     # Save chunks to a JSON file
#     with open(output_file, "w", encoding="utf-8") as f:
#         json.dump(chunks, f, indent=4, ensure_ascii=False)

#     print(f"Chunks saved successfully to {output_file}")


import os
import json
import re
import nltk
from nltk.tokenize import sent_tokenize
from PyPDF2 import PdfReader
from docx import Document

nltk.download('punkt')

class ChunkingService:
    def __init__(self, chunk_size=200, overlap=50):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def extract_text(self, file_path):
        text = ""
        try:
            if file_path.lower().endswith(".pdf"):
                with open(file_path, "rb") as file:
                    reader = PdfReader(file)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            elif file_path.lower().endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as file:
                    text = file.read().strip()

            elif file_path.lower().endswith(".docx"):
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs]).strip()

            else:
                raise ValueError("Unsupported file format. Only PDF, TXT, and DOCX are allowed.")

        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
        
        return text

    def clean_text(self, text):
        text = re.sub(r'\s+', ' ', text)
        text = text.replace("ﬁ", "fi").replace("ﬂ", "fl")  # Fix encoding issues
        return text.strip()

    def split_into_sentences(self, text):
        return sent_tokenize(text)

    def create_chunks(self, sentences):
        chunks = []
        current_chunk = []

        for sentence in sentences:
            sentence = sentence.strip()
            if sum(len(s) for s in current_chunk) + len(sentence) <= self.chunk_size:
                current_chunk.append(sentence)
            else:
                chunks.append(" ".join(current_chunk))
                current_chunk = [sentence]

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        # Apply overlap
        overlapped_chunks = []
        for i in range(len(chunks)):
            start = max(0, i - 1)  # Include previous chunk for overlap
            overlap_text = " ".join(chunks[start:i + 1])
            overlapped_chunks.append(overlap_text[:self.chunk_size].strip())

        return overlapped_chunks

    def process(self, pdf_path):
        raw_text = self.extract_text(pdf_path)
        cleaned_text = self.clean_text(raw_text)
        sentences = self.split_into_sentences(cleaned_text)
        chunk_texts = self.create_chunks(sentences)

        return [{"chunk_number": i + 1, "text": chunk} for i, chunk in enumerate(chunk_texts)]


# if __name__ == "__main__":
#     chunker = ChunkingService(chunk_size=300, overlap=50)
#     file_path = r"C:/Users/srees/OneDrive/Desktop/Semester-08/finalNotes.pdf"
#     chunks = chunker.process(file_path)

#     # Convert chunks into a list of dictionaries
#     chunk_dicts = [{"chunk_number": i + 1, "text": chunk} for i, chunk in enumerate(chunks)]

#     # Save chunks to JSON file
#     output_dir = r"C:\Users\srees\OneDrive\Desktop\Semester-08\capstone2-impl\app"
#     os.makedirs(output_dir, exist_ok=True)
#     output_file = os.path.join(output_dir, os.path.basename(file_path) + "_chunks.json")

#     with open(output_file, "w", encoding="utf-8") as f:
#         json.dump(chunk_dicts, f, indent=4)

#     print(f"Chunks saved to {output_file}")

#     # Print chunks for verification
#     for chunk in chunk_dicts:
#         print(f"Chunk {chunk['chunk_number']}:\n{chunk['text']}\n")

